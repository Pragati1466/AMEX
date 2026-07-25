"""
Document Processing & OCR Module.
Handles extraction of raw text from uploaded evidence documents.
Supports: PDF (PyMuPDF), PDF Tables (pdfplumber), Word (python-docx), Images (Tesseract OCR).
"""

import io
import os
import tempfile
from pathlib import Path
from typing import Optional

from loguru import logger

# PDF
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    logger.warning("PyMuPDF not available. PDF text extraction disabled.")

# PDF Tables
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False
    logger.warning("pdfplumber not available. PDF table extraction disabled.")

# Word
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not available. Word document extraction disabled.")

# Image OCR
try:
    import pytesseract
    from PIL import Image
    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False
    logger.warning("pytesseract not available. Image OCR disabled.")


# --------------- Supported File Types ---------------

PDF_MIME_TYPES = {"application/pdf"}
PDF_EXTENSIONS = {".pdf"}

WORD_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}
WORD_EXTENSIONS = {".docx", ".doc"}

IMAGE_MIME_TYPES = {
    "image/png",
    "image/jpeg",
    "image/jpg",
    "image/tiff",
    "image/bmp",
    "image/webp",
}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp"}

TEXT_EXTENSIONS = {".txt", ".csv", ".json", ".xml", ".html", ".md"}


# --------------- Public API ---------------

def extract_text_from_bytes(
    file_data: bytes,
    filename: str,
    mime_type: Optional[str] = None,
) -> str:
    """
    Extract raw text from an uploaded document based on its file type.

    Args:
        file_data: Raw bytes of the uploaded file.
        filename: Original filename (used to determine file type).
        mime_type: Optional MIME type (used as additional hint).

    Returns:
        Extracted raw text content. Empty string if extraction fails or
        the file type is unsupported.
    """
    ext = _get_extension(filename).lower()
    mime = (mime_type or "").lower()

    # Determine file category
    if ext in PDF_EXTENSIONS or mime in PDF_MIME_TYPES:
        return _extract_pdf_text(file_data)

    if ext in WORD_EXTENSIONS or mime in WORD_MIME_TYPES:
        return _extract_word_text(file_data)

    if ext in IMAGE_EXTENSIONS or mime in IMAGE_MIME_TYPES:
        return _extract_image_text(file_data)

    if ext in TEXT_EXTENSIONS:
        return _extract_plain_text(file_data)

    # Fallback: try OCR on unknown binary files
    if _looks_like_image(file_data):
        logger.info(f"Attempting OCR on unknown file type: {filename}")
        return _extract_image_text(file_data)

    # Last resort: try to decode as text
    return _extract_plain_text(file_data)


def get_supported_extensions() -> set:
    """Return all supported file extensions."""
    return PDF_EXTENSIONS | WORD_EXTENSIONS | IMAGE_EXTENSIONS | TEXT_EXTENSIONS


def is_supported_file(filename: str) -> bool:
    """Check if a file is supported for text extraction."""
    ext = _get_extension(filename).lower()
    return ext in get_supported_extensions()


# --------------- PDF Extraction ---------------

def _extract_pdf_text(file_data: bytes) -> str:
    """
    Extract text from a PDF using PyMuPDF.
    Also attempts table extraction with pdfplumber as a supplement.
    """
    parts = []

    # 1. PyMuPDF text extraction (best for general PDF text)
    if HAS_PYMUPDF:
        try:
            doc = fitz.open(stream=file_data, filetype="pdf")
            text_parts = []
            for page_num, page in enumerate(doc, start=1):
                page_text = page.get_text().strip()
                if page_text:
                    text_parts.append(f"--- Page {page_num} ---\n{page_text}")
            doc.close()
            if text_parts:
                parts.append("\n\n".join(text_parts))
                logger.debug(f"Extracted {sum(len(t) for t in text_parts)} chars from PDF via PyMuPDF")
        except Exception as e:
            logger.error(f"PyMuPDF extraction failed: {e}")

    # 2. pdfplumber table extraction (supplements with tabular data)
    if HAS_PDFPLUMBER:
        try:
            with pdfplumber.open(io.BytesIO(file_data)) as pdf:
                table_parts = []
                for page_num, page in enumerate(pdf.pages, start=1):
                    tables = page.extract_tables()
                    if tables:
                        table_texts = []
                        for table_idx, table in enumerate(tables, start=1):
                            if table:
                                rows = []
                                for row in table:
                                    cleaned = [str(cell).strip() if cell else "" for cell in row]
                                    rows.append(" | ".join(cleaned))
                                table_texts.append(
                                    f"[Table {table_idx} on Page {page_num}]\n" + "\n".join(rows)
                                )
                        if table_texts:
                            table_parts.append("\n\n".join(table_texts))
                if table_parts:
                    parts.append("\n\n".join(table_parts))
                    logger.debug(f"Extracted {len(table_parts)} tables from PDF via pdfplumber")
        except Exception as e:
            logger.error(f"pdfplumber extraction failed: {e}")

    if not parts:
        logger.warning("No text extracted from PDF")
        return ""

    return "\n\n".join(parts)


# --------------- Word Document Extraction ---------------

def _extract_word_text(file_data: bytes) -> str:
    """Extract text from a Word document (.docx) using python-docx."""
    if not HAS_DOCX:
        logger.error("python-docx not installed. Cannot extract Word document.")
        return ""

    try:
        doc = docx.Document(io.BytesIO(file_data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract tables
        table_texts = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_texts.append("\n".join(rows))

        parts = []
        if paragraphs:
            parts.append("\n\n".join(paragraphs))
        if table_texts:
            parts.append("\n\n".join(table_texts))

        result = "\n\n".join(parts)
        logger.debug(f"Extracted {len(result)} chars from Word document")
        return result

    except Exception as e:
        logger.error(f"Word document extraction failed: {e}")
        return ""


# --------------- Image OCR ---------------

def _extract_image_text(file_data: bytes) -> str:
    """Extract text from an image using Tesseract OCR."""
    if not HAS_TESSERACT:
        logger.error("pytesseract not installed. Cannot perform OCR.")
        return ""

    try:
        # Check if tesseract executable is available
        try:
            pytesseract.get_tesseract_version()
        except Exception:
            logger.error("Tesseract OCR engine not found on system. Install with: brew install tesseract")
            return ""

        image = Image.open(io.BytesIO(file_data))

        # Convert to RGB if necessary (e.g., RGBA PNG)
        if image.mode not in ("L", "RGB"):
            image = image.convert("RGB")

        # Perform OCR
        text = pytesseract.image_to_string(image, lang="eng")
        text = text.strip()

        logger.debug(f"Extracted {len(text)} chars from image via OCR")
        return text

    except Exception as e:
        logger.error(f"Image OCR failed: {e}")
        return ""


# --------------- Plain Text ---------------

def _extract_plain_text(file_data: bytes) -> str:
    """Decode raw bytes as UTF-8 text."""
    try:
        return file_data.decode("utf-8").strip()
    except UnicodeDecodeError:
        try:
            return file_data.decode("latin-1").strip()
        except Exception as e:
            logger.error(f"Plain text decoding failed: {e}")
            return ""


# --------------- Helpers ---------------

def _get_extension(filename: str) -> str:
    """Get the file extension from a filename."""
    return Path(filename).suffix.lower()


def _looks_like_image(file_data: bytes) -> bool:
    """Check if the file data starts with a known image magic byte sequence."""
    if len(file_data) < 8:
        return False
    # PNG
    if file_data[:8] == b"\x89PNG\r\n\x1a\n":
        return True
    # JPEG
    if file_data[:2] == b"\xff\xd8":
        return True
    # BMP
    if file_data[:2] == b"BM":
        return True
    # TIFF (little-endian)
    if file_data[:2] in (b"II", b"MM"):
        return True
    # WebP
    if file_data[:4] == b"RIFF" and file_data[8:12] == b"WEBP":
        return True
    return False